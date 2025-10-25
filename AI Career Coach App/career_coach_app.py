import gradio as gr
import os
import json 
import re 
from pypdf import PdfReader
import pandas as pd 
import numpy as np
import copy
from docx import Document as DocxDocument 
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import atexit 
import shutil 
from datetime import datetime

# --- IBM watsonx.ai Initialization (Ensure your credentials are here) ---
from ibm_watsonx_ai.foundation_models import ModelInference
from ibm_watsonx_ai import Credentials
from ibm_watsonx_ai.foundation_models.schema import TextChatParameters

# NOTE: Replace with your actual credentials and project ID
credentials = Credentials(url = "https://us-south.ml.cloud.ibm.com") 
model_id = "ibm/granite-3-8b-instruct"  
project_id = "skills-network"  # Replace with your actual project ID
base_params = TextChatParameters(temperature=0.3, max_tokens=2048)
model = ModelInference(model_id=model_id, credentials=credentials, project_id=project_id, params=base_params)

# Global list to track temporary files for cleanup
temp_files_to_clean = []

def cleanup_temp_files():
    """Removes all generated temporary files when the script exits."""
    for file_path in temp_files_to_clean:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Error cleaning up file {file_path}: {e}")

# Register the cleanup function to run at exit
atexit.register(cleanup_temp_files)


def llm_chat(prompt, max_tokens=None):
    """Generic function to send a prompt to the initialized model."""
    messages = [{"role": "user", "content": [{"type": "text", "text": prompt}]}] 
    
    current_params = copy.copy(base_params)
    if max_tokens:
        current_params.max_tokens = max_tokens
        current_model = ModelInference(model_id=model_id, credentials=credentials, project_id=project_id, params=current_params)
    else:
        current_model = model

    try:
        generated_response = current_model.chat(messages=messages)
        generated_text = generated_response['choices'][0]['message']['content']
        return generated_text.strip()
    except Exception as e:
        return f"Error communicating with LLM: {e}. Check watsonx.ai setup (credentials, project ID, region, model_id)."

# --------------------------------------------------------------------------
# --- CORE HELPER FUNCTIONS ---
# --------------------------------------------------------------------------

def extract_text_from_file(file_input):
    """Handles file extraction for PDF, TXT, and MD, cleaning Markdown syntax."""
    if file_input is None:
        return "Error: No file uploaded."
        
    file_path = file_input.name
    
    try:
        if file_path.lower().endswith('.pdf'):
            reader = PdfReader(file_path)
            text = "".join(page.extract_text() + "\n\n" if page.extract_text() else "" for page in reader.pages)
        
        # NEW: Check for both .txt and .md
        elif file_path.lower().endswith(('.txt', '.md')):
            
            # Read the raw text content
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # CRITICAL CLEANUP: If it's a Markdown file, strip the syntax
            if file_path.lower().endswith('.md'):
                text = clean_markdown_input(text)
        
        else:
            return "Error: File type not supported for text extraction (only PDF/TXT/MD)."

        if not text.strip():
            return "Error: Text extraction failed or document is empty."
        
        return text.strip()
    except Exception as e:
        return f"File Extraction Failed: {e}"

def clean_markdown_input(text_content):
    """
    Cleans text content by removing common Markdown formatting characters.
    It removes: headers (#), bold (**), italics (*), list markers (-), 
    and converts [text](link) to just 'text'.
    """
    if not text_content:
        return ""
    
    # 1. Remove Markdown links: [text](link) becomes text
    text_content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text_content)
    
    # 2. Remove common inline and block formatting: **, *, #, >, -, +, =
    # This addresses bolding, italics, headers, blockquotes, and list items
    text_content = re.sub(r'(\*\*|__|#|\*|~|`|>|-|\+|=|\n\s*-)', r' ', text_content)
    
    # 3. Clean up excessive whitespace created by removal
    text_content = re.sub(r'\s+', ' ', text_content).strip()
    
    return text_content

def enrich_skills_from_certifications(current_skills_text, education_certifications_text):
    """
    Uses the LLM to extract technical skills implied by certifications 
    and adds them to the existing skills list, avoiding duplicates.
    """
    if not education_certifications_text or education_certifications_text.strip() == "":
        return current_skills_text
    
    # 1. LLM Prompt to Extract New Skills
    prompt = f"""
    Analyze the provided Certifications and Education text. Your task is to extract a list of **specific, technical skills** (e.g., Python, Scikit-learn, MLOps, NLP, Time Series Analysis) that are strongly implied or explicitly taught by these programs.
    
    CRITICAL RULE: Ignore general soft skills or degrees (e.g., Bachelor's degree, Leadership). Only list specific, valuable technical skills.
    
    Certifications and Education Text: '''{education_certifications_text}'''
    
    OUTPUT RULE: Return ONLY a comma-separated list of technical skills. Do not include any preambles, titles, or descriptions.
    """
    
    try:
        # Use a small token limit since we only want a list
        extracted_skills_raw = llm_chat(prompt, max_tokens=512).strip()
    except Exception as e:
        # Fail gracefully if LLM call fails
        print(f"LLM skill extraction failed: {e}")
        return current_skills_text

    # 2. Process and Deduplicate Skills
    
    # Clean and split the extracted skills
    extracted_skills_list = [
        re.sub(r'[^a-zA-Z0-9\s\-\.\#\+]', '', s).strip() # Clean punctuation
        for s in extracted_skills_raw.split(',')
        if s.strip()
    ]
    
    # Process current skills list
    current_skills_list = [
        re.sub(r'[^a-zA-Z0-9\s\-\.\#\+]', '', s).strip()
        for s in current_skills_text.split('\n') 
        if s.strip() and not s.lower().startswith("error")
    ]
    
    # Combine and deduplicate (case-insensitive check for uniqueness)
    unique_skills_set = {s.lower(): s for s in current_skills_list}
    
    for skill in extracted_skills_list:
        if skill.lower() not in unique_skills_set:
            unique_skills_set[skill.lower()] = skill
            
    # Reconstruct the skills text, keeping the original formatting style (one skill per line)
    updated_skills_text = "\n".join(unique_skills_set.values())
    
    return updated_skills_text

def clean_structured_data(data):
    """Converts nested lists/dicts into a clean, multi-line string for Gradio textboxes."""
    if isinstance(data, list):
        clean_lines = []
        for item in data:
            if isinstance(item, dict):
                lines = [f"**{str(k).upper()}**: {str(v)}" for k, v in item.items() if v]
                clean_lines.append("\n".join(lines).strip())
                clean_lines.append("-" * 30) 
            else:
                clean_lines.append(str(item).strip())
        
        while clean_lines and clean_lines[-1] == "-" * 30:
            clean_lines.pop()
            
        return "\n".join(clean_lines).strip()
    elif isinstance(data, dict):
        lines = [f"**{str(k).upper()}**: {str(v)}" for k, v in data.items() if v]
        return "\n".join(lines).strip()
    return str(data).strip()

def extract_name(contact_info):
    """
    Extracts the candidate's name from the structured contact information string 
    (e.g., from **NAME**: Gregory Charles).
    """
    if not contact_info:
        return "Candidate"
    
    # Regex to find '**NAME**:' followed by any characters until a newline or end of string
    match = re.search(r'\*\*NAME\*\*:\s*(.*?)\s*(?:\n|$)', contact_info, re.IGNORECASE)
    
    if match:
        name = match.group(1).strip()
        return name.replace('*', '').strip()
    
    return "Candidate"

def save_docx(content, filename_base):
    """Creates and saves a DOCX file from text content."""
    # NOTE: Requires docx imports, os, and re
    try:
        document = DocxDocument()
        
        # Simple font size/style setup
        style = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        for line in content.split('\n'):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Simple bolding for headings (improves formatting from LLM plain text)
            if line.isupper() and len(line.split()) < 4 and line.strip():
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(14) # Larger font for section titles
            elif line.startswith('-'):
                # Add bullet points for list items
                p.style = 'List Bullet'
                p.add_run(line[1:].strip())
            else:
                p.add_run(line)

        # Ensure safe directory creation and file saving
        if not os.path.exists('temp_files'):
            os.makedirs('temp_files', exist_ok=True)
            
        safe_filename = re.sub(r'[^\w\-_\.]', '_', filename_base)
        temp_filepath = os.path.join(os.getcwd(), 'temp_files', safe_filename)
        
        # Ensure temp_files_to_clean is defined globally at the top of your script
        # temp_files_to_clean.append(temp_filepath) # Uncomment if you define this list globally
        
        document.save(temp_filepath)
        return temp_filepath
    except Exception as e:
        return f"Error creating DOCX file: {e}"

def safe_json_parse(raw_llm_output):
    """Attempts to extract and parse a JSON object from an LLM string."""
    json_string = raw_llm_output.strip()

    if json_string.startswith(("```json", "```")):
        json_string = re.sub(r'^(```json|```)\s*|\s*(```)$', '', json_string, flags=re.IGNORECASE | re.MULTILINE).strip()
    
    json_match = re.search(r'\{[\s\S]*\}', json_string)
    
    if json_match:
        json_string = json_match.group(0).strip()
        
        try:
            return json.loads(json_string), None
        except json.JSONDecodeError as e:
            json_string_fixed = json_string.replace('...', '').replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
            try:
                return json.loads(json_string_fixed), None
            except json.JSONDecodeError:
                return None, f"JSON Decode Error (Attempted Fix): {e}"
    
    return None, f"JSON Structure Error: No valid JSON object found in output. Raw: {json_string[:500]}..."

# --------------------------------------------------------------------------
# --- CLASSIFICATION AND DATE HELPERS ---
# --------------------------------------------------------------------------

def parse_date_for_sort(date_str):
    """Converts date string to datetime object for sorting. Returns FUTURE_DATE for 'Present' or NaT."""
    FUTURE_DATE = pd.to_datetime('2099-12-31') 
    if pd.isna(date_str) or str(date_str).strip().lower() in ('nan', 'present', ''):
        return FUTURE_DATE
    try:
        # Attempt to parse common formats
        return pd.to_datetime(date_str, errors='coerce', dayfirst=False)
    except Exception:
        return pd.NaT

# Modified to distinguish Academic Degrees from Certifications
def classify_education_entry(row):
    """Assigns a category and sort priority to an education or certification entry."""
    name_lower = str(row.get('Name', row.get('Degree Name', ''))).lower()
    
    # 1. ACADEMIC DEGREES (Diplomas) - High Precedence
    if 'phd' in name_lower or 'doctorate' in name_lower:
        return "Academic Degrees (PhD/Doctorate)", 3 
    if 'master' in name_lower or 'm.sc.' in name_lower or 'm.a.' in name_lower:
        return "Academic Degrees (Master's)", 2
    if 'bachelor' in name_lower or 'b.sc.' in name_lower or 'b.a.' in name_lower or 'high school diploma' in name_lower:
        return "Academic Degrees (Bachelor's & High School)", 1
    
    # 2. PROFESSIONAL CERTIFICATIONS - Lower Precedence (sorted by internal category)
    # Data Science/ML/AI
    if any(k in name_lower for k in ['data science', 'machine learning', 'mlops', 'python', 'sql', 'tableau', 'generative ai', 'cs50', 'data analytics', 'mit']):
        return "Professional Certifications (Data Science/ML)", 0
    # Teaching/Language
    if any(k in name_lower for k in ['teaching english', 'tefl', 'young learners', 'business english', 'tquk', 'foreign language', 'english teacher', 'language school', 'language institute']):
        return "Professional Certifications (Teaching/Language)", 0
    # Automotive/Engineering
    if any(k in name_lower for k in ['automotive', 'varta', 'mercedes-benz', 'hev', 'hybrid', 'electric vehicles', 'engineering', 'product qualification', 'electricity', 'technician']):
        return "Professional Certifications (Engineering)", 0
    # Business/Leadership/Strategy
    if any(k in name_lower for k in ['business management', 'negotiation', 'leadership', 'strategic thinking', 'product management', 'business analytics', 'wharton', 'leading people', 'interviewing', 'career planning']):
        return "Professional Certifications (Business/Leadership)", 0
    # Others
    return "Professional Certifications (Other)", 0

# --------------------------------------------------------------------------
# --- MD FILE LOADER ---
# --------------------------------------------------------------------------
def parse_md_file(file_input):
    """Reads a master profile MD file and splits it into the six profile fields, cleaning internal Markdown."""
    if file_input is None:
        return "", "", "", "", "", "", "Error: No file uploaded."
        
    file_path = file_input.name
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        DELIMITER_MAP = {
            "CONTACT:": "CONTACT", 
            "SUMMARY:": "SUMMARY", 
            "EXPERIENCE (Detailed Career History & Achievements):": "EXPERIENCE", 
            "PROJECTS (Key personal/work projects, Awards, Publications):": "PROJECTS", 
            "EDUCATION (Degrees and Certifications):": "EDUCATION", 
            "SKILLS (All Technical and Soft Skills):": "SKILLS"
        }
        
        sections = {key: "" for key in DELIMITER_MAP.values()}
        
        # Initial cleanup of file boundary markers
        content = content.replace("--- CANDIDATE MASTER PROFILE START (Comprehensive Data) ---\n", "").strip()
        
        current_section = None
        
        lines = content.split('\n')
        content_start_index = next((i for i, line in enumerate(lines) if line.startswith(tuple(DELIMITER_MAP.keys()))), 0)
        
        for line in lines[content_start_index:]:
            line_stripped = line.strip()
            
            found_key = next((key for key in DELIMITER_MAP.keys() if line_stripped.startswith(key)), None)
            
            if found_key:
                current_section = DELIMITER_MAP[found_key]
            elif current_section:
                if line_stripped and line_stripped != found_key:
                    if sections[current_section]:
                        sections[current_section] += "\n" + line
                    else:
                        sections[current_section] += line
                
        
        # --- Aggressive Cleanup of internal delimiters ---
        
        # Pattern to match the merge marker with any surrounding whitespace/newlines
        MERGE_MARKER_PATTERN = r'\s*--- NEW [A-Z/ &]+ ---\s*' 
        
        # General separators used for experience blocks and file end marker
        NOISE_PATTERNS = [
            r'\n-{30,}\n',   # Experience Separator (----------------------------------------)
            r'\s*--- CANDIDATE MASTER PROFILE END ---\s*' # Final file boundary marker
        ]
        
        for key in sections:
            if sections[key]: 
                content = sections[key]
                
                # 1. REMOVE ALL MERGE MARKERS 
                content = re.sub(MERGE_MARKER_PATTERN, '\n\n', content, flags=re.IGNORECASE).strip()
                
                # 2. REMOVE OTHER GENERAL NOISE PATTERNS
                for pattern in NOISE_PATTERNS:
                    content = re.sub(pattern, '\n', content, flags=re.MULTILINE | re.IGNORECASE).strip() 
                
                # 3. REMOVE INTERNAL MARKDOWN CLUTTER 
                content = content.replace('**', '') 
                content = content.replace('### ', '') 
                
                # --- SECTION-SPECIFIC FORMATTING ---
                
                if key == 'CONTACT' or key == 'SUMMARY':
                    # Ensures single-line values stay on the same line as the header (e.g., NAME: Gregory Charles)
                    content = re.sub(r'([A-Z ]+[:])\s*\n\s*', r'\1 ', content, flags=re.MULTILINE).strip()

                elif key == 'EDUCATION' or key == 'SKILLS':
                    # Education/Skills Readability Fix (Header: \n\n- Item)

                    # 4a. Ensure there are AT LEAST two newlines before any new group header for separation
                    # This rule looks for content (a list item) followed by a newline and then a header, inserting separation.
                    content = re.sub(r'([^\n])\n([A-Z/ &]+[:])', r'\1\n\n\2', content, flags=re.MULTILINE).strip()

                    # 4b. Ensure a blank line separates the group header from the list items below them
                    content = re.sub(r'([A-Z/ &]+[:])\n', r'\1\n\n', content, flags=re.MULTILINE).strip()

                    # 4c. Ensure items within the same list are only separated by a single newline
                    content = re.sub(r'(\n- [^\n]+)\n\n(- [^\n]+)', r'\1\n\2', content, flags=re.MULTILINE)
                    
                    if key == 'SKILLS':
                        # Fix for skills list where items are comma-separated and should be kept on one line under the header
                        content = re.sub(r'([A-Z/ &]+[:])\s*\n\n\s*([A-Z])', r'\1 \2', content, flags=re.MULTILINE)
                        

                elif key == 'PROJECTS':
                    # 5a. Project Separator Fix: Replace the '========================================' with '-----'
                    content = re.sub(r'\n={30,}\n', '\n-----\n\n', content, flags=re.MULTILINE)
                    
                    # 5b. Project Block Separation: Ensure a blank line before every new project header
                    content = re.sub(r'(\nPROJECT/AWARD:)', r'\n\n\1', content, flags=re.MULTILINE).strip()

                    # 5c. Fix date splits (e.g., (Dec 2024\n - Dec 2024))
                    content = re.sub(r'(\(\S+ \d{4})\s*\n\s*-\s*(\S+ \d{4}\))', r'\1 - \2', content)

                    # 5d. Ensure DATES is followed by a blank line before DETAILS
                    content = re.sub(r'(DATES:.*?\))\n(DETAILS:)', r'\1\n\n\2', content, flags=re.MULTILINE | re.DOTALL)
                    
                    # 5e. Project Details Fix: Ensure DETAILS: is on its own line followed by a blank line before content
                    content = re.sub(r'(DETAILS:)\s*(.*)', r'\1\n\n\2', content, flags=re.DOTALL)
                    
                    # 5f. Break dense text into list items: Add a newline before any hyphen that is not preceded by a newline/whitespace
                    content = re.sub(r'([a-z\)\.])\s*-\s+([A-Z])', r'\1\n\n- \2', content, flags=re.MULTILINE)
                    
                    # 5g. Clean up list item spacing to ensure single newline between list items
                    content = re.sub(r'(\n- [^\n]+)\n\n(- [^\n]+)', r'\1\n\2', content, flags=re.MULTILINE)
                    
                # 6. Final cleaning of excessive blank lines (converts 3+ newlines into 2)
                content = re.sub(r'\n{3,}', '\n\n', content).strip()
                
                sections[key] = content
        # ----------------------------------------------------------------------------------

        # Final cleanup and return
        return (
            sections.get('CONTACT', '').strip(),
            sections.get('SUMMARY', '').strip(),
            sections.get('EXPERIENCE', '').strip(), 
            sections.get('EDUCATION', '').strip(), 
            sections.get('SKILLS', '').strip(), 
            sections.get('PROJECTS', '').strip(), 
            "✅ Master Profile (.md) loaded successfully into profile fields and cleaned of internal merging markers."
        )

    except Exception as e:
        return "", "", "", "", "", "", f"MD File Loading Failed: {e}"                           
# --------------------------------------------------------------------------
# --- CORE LOGIC FUNCTIONS ---
# --------------------------------------------------------------------------

def parse_and_distribute_file(file_input):
    """Main file handler that decides between PDF/TXT parsing (via LLM), CSV processing, or MD loading."""
    # Return 6 textboxes, status, state, 2 radios (interactive), 1 button (interactive)
    default_outputs = [gr.update()] * 6 + [gr.update()] + [None] + [gr.update(interactive=False)] * 2 + [gr.update(interactive=False)]

    if file_input is None:
        default_outputs[6] = "Error: No file uploaded."
        return default_outputs
        
    file_path = file_input.name
    
    if file_path.lower().endswith('.csv'):
        return csv_dispatcher(file_path)
    elif file_path.lower().endswith(('.pdf', '.txt')):
        # PDF/TXT processing bypasses the new UI controls
        parsed_outputs = parse_and_distribute_pdf_txt(file_input)
        return parsed_outputs + [None] + [gr.update(interactive=False)] * 2 + [gr.update(interactive=False)]
    elif file_path.lower().endswith('.md'): # <-- NEW ADDITION for MD loading
        # MD parsing bypasses the new UI controls (it overwrites all boxes)
        parsed_outputs = parse_md_file(file_input)
        # parsed_outputs is a tuple of (6 boxes, status)
        return list(parsed_outputs[:6]) + [parsed_outputs[6]] + [None] + [gr.update(interactive=False)] * 2 + [gr.update(interactive=False)]
    else:
        default_outputs[6] = "Error: Unsupported file type. Please upload PDF, TXT, CSV, or MD."
        return default_outputs


def parse_and_distribute_pdf_txt(file_input):
    """Parses raw resume text (PDF/TXT) into structured fields using LLM."""
    raw_text = extract_text_from_file(file_input)
    if raw_text.startswith("Error"):
        return "", "", "", "", "", "", raw_text 

    parsing_prompt = f"""
    Analyze the raw resume text provided below and segment it into six distinct sections: CONTACT_INFO, SUMMARY, EXPERIENCE, EDUCATION, SKILLS, and PROJECTS.
    
    CRITICAL RULE: You MUST return the results STRICTLY as a single JSON object. Do NOT include ANY text, comments, markdown wrappers (e.g., ```json), or explanations outside of the single JSON object.
    
    JSON Keys must be exactly: "CONTACT_INFO", "SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS".
    
    Raw Resume Text:
    ---
    {raw_text[:4096]} 
    ---
    """
    
    try:
        raw_llm_output = llm_chat(parsing_prompt, max_tokens=2048) 
        
        sections, error = safe_json_parse(raw_llm_output)
        
        if error:
            return "", "", "", "", "", "", f"Parsing Error: {error}"
            
        return (
            clean_structured_data(sections.get('CONTACT_INFO', '')), 
            clean_structured_data(sections.get('SUMMARY', '')), 
            clean_structured_data(sections.get('EXPERIENCE', '')), 
            clean_structured_data(sections.get('EDUCATION', '')), 
            clean_structured_data(sections.get('SKILLS', '')), 
            clean_structured_data(sections.get('PROJECTS', '')), 
            "✅ PDF/TXT Parsed successfully into profile fields using LLM. Review and click 'Save'."
        )

    except Exception as e:
        return "", "", "", "", "", "", f"Parsing Error: An unexpected error occurred: {e}."


def csv_dispatcher(csv_file_path):
    """Checks CSV by name first, then by columns. Prompts user for action if ambiguous."""
    # Returns: 6 textboxes updates, status, file_state, radio_destination, radio_action, button
    updates = [gr.update()] * 6 + [gr.update()] + [None] + [gr.update(interactive=False)] * 3
    file_name = os.path.basename(csv_file_path).lower()
    
    # --- 1. IMMEDIATE PROCESSING (Profile.csv is always processed immediately) ---
    if "profile" in file_name:
        formatted_content = _format_csv_data(csv_file_path, "Profile")
        if isinstance(formatted_content, str) and formatted_content.startswith("Error"):
             updates[6] = formatted_content
             return updates
             
        updates[0] = gr.update(value=formatted_content.get('CONTACT', ''))
        updates[1] = gr.update(value=formatted_content.get('SUMMARY', ''))
        updates[6] = f"✅ CSV File **{file_name}** automatically loaded to **Contact & Summary**."
        return updates

    # --- 2. AMBIGUOUS/NON-PROFILE FILES (Requires User Choice) ---
    # Determine the most likely default, but enforce user choice for merge/overwrite
    
    default_destination = "Experience"
    if "certif" in file_name or "educat" in file_name:
        default_destination = "Education"
    elif "project" in file_name or "award" in file_name:
        default_destination = "Projects"
    elif "skill" in file_name:
        default_destination = "Skills"
    elif "positions" in file_name:
        default_destination = "Experience"
    
    status = (
        f"❓ CSV File **{file_name}** uploaded. It appears to be for the **{default_destination}** section. "
        f"Please **confirm the Destination Section** and choose to **Overwrite** or **Append** before clicking 'Process Selected CSV'."
    )
    
    # Enable and set default values for the new controls
    updates[7] = csv_file_path  # Set CSV_FILE_TO_PROCESS state
    updates[6] = status
    updates[8] = gr.update(value=default_destination, interactive=True)
    updates[9] = gr.update(interactive=True)
    updates[10] = gr.update(interactive=True)
    
    return updates


def _format_csv_data(csv_file_path, destination):
    """Formats the CSV data into a professional string based on the target section, with sorting/grouping."""
    try:
        df = pd.read_csv(csv_file_path)
    except Exception as e:
        return f"Error: Could not read CSV file at {csv_file_path}. {e}"
        
    # Standardize NaN handling for strings
    df = df.replace({np.nan: ''})

    # --- Profile/Personal Info Format ---
    if destination == "Profile":
        # ... [Existing Profile logic remains the same] ...
        if 'First Name' not in df.columns or 'Summary' not in df.columns:
            return "Error: Profile CSV missing 'First Name' or 'Summary' column."
            
        if df.empty: return {'CONTACT': '', 'SUMMARY': ''}
            
        row = df.iloc[0]
        
        contact_details = {
            "Name": f"{row.get('First Name', '')} {row.get('Last Name', '')}".strip(),
            "Headline": row.get('Headline', ''),
            "Address": row.get('Address', ''),
            "Location": row.get('Geo Location', ''),
            "Industry": row.get('Industry', ''),
            "Websites": row.get('Websites', ''),
            "Twitter": row.get('Twitter Handles', ''),
        }
        contact = clean_structured_data({k: v for k, v in contact_details.items() if v and str(v).strip().lower() != 'nan'})
        summary = str(row.get('Summary', '')).strip()
        return {'CONTACT': contact, 'SUMMARY': summary}


    # --- Experience/Work History Format (NOW SORTED BY DATE) ---
    elif destination == "Experience":
        output_text = ["--- NEW EXPERIENCE ENTRIES ---"]
        
        # Robust column fetching for Experience
        title_col = 'Title' if 'Title' in df.columns else 'Position Name'
        company_col = 'Company Name' if 'Company Name' in df.columns else 'Company'
        start_col = 'Started On' if 'Started On' in df.columns else 'Start Date'
        finish_col = 'Finished On' if 'Finished On' in df.columns else 'End Date'

        # 1. Prepare for Sorting: Create sort key columns
        df['End_Date_Obj'] = df[finish_col].apply(parse_date_for_sort)
        df['Start_Date_Obj'] = df[start_col].apply(parse_date_for_sort)
        
        # 2. Sort: Reverse Chronological (Most recent first).
        # We use End_Date_Obj (which has FUTURE_DATE for 'Present') and then Start_Date_Obj for ties.
        df = df.sort_values(by=['End_Date_Obj', 'Start_Date_Obj'], ascending=[False, False])
        
        for _, row in df.iterrows():
            title = row.get(title_col, '')
            company = row.get(company_col, '')
            
            # Date Handling
            start = str(row.get(start_col, '')).split('T')[0].strip()
            end_raw = row.get(finish_col)
            end = str(end_raw).split('T')[0].strip()
            end_display = end if end and end.lower() not in ('nan', 'present') else 'Present'
            
            description = str(row.get('Description', '')).replace('\n', ' ')
            
            if not title.strip(): continue # Skip empty entries
            
            output_text.append(f"TITLE: **{title}**")
            output_text.append(f"COMPANY: {company} ({start} - {end_display})")
            if description and description.lower() != 'nan':
                output_text.append(f"DETAILS: {description}")
            output_text.append("-" * 40)
            
        while output_text and output_text[-1] == "-" * 40:
            output_text.pop()
            
        return "\n".join(output_text).strip()

    # --- Projects/Awards Format (NEW LOGIC HERE) ---
    elif destination == "Projects":
        output_text = ["--- NEW PROJECTS/AWARDS ---"]
        
        title_col = 'Name' if 'Name' in df.columns else 'Project Name'
        
        for _, row in df.iterrows():
            title = row.get(title_col, row.get('Title', ''))
            
            start_raw = row.get('Start Date', row.get('Started On', ''))
            end_raw = row.get('End Date', row.get('Finished On', ''))
            
            start = str(start_raw).split('T')[0].strip()
            end = str(end_raw).split('T')[0].strip()
            end_display = end if end and end.lower() not in ('nan', 'present') else '' 

            description = str(row.get('Description', '')).replace('\n', ' ')
            url = str(row.get('Url', row.get('Project Url', '')))
            
            if not title.strip(): continue

            # --- Differentiation Heuristic ---
            entry_type = "PROJECT"
            title_lower = title.lower()
            
            # Simple keyword check for Awards/Publications
            award_keywords = ['award', 'nominee', 'publication', 'patent', 'recognized', 'scholarship', 'prize', 'honor']
            if any(kw in title_lower for kw in award_keywords) or any(kw in description.lower() for kw in award_keywords):
                entry_type = "AWARD/RECOGNITION"
            # END Differentiation Heuristic

            output_text.append(f"{entry_type}: **{title}**")
            
            dates_str = f" ({start}{' - ' + end_display if end_display else ''})"
            if dates_str.strip() not in ('()', '( - '):
                output_text.append(f"DATES: {dates_str.strip()}")
                
            if description and description.lower() != 'nan':
                output_text.append(f"DETAILS: {description}")
            if url and url.lower() != 'nan':
                 output_text.append(f"LINK: {url}")
            output_text.append("=" * 40)
            
        while output_text and output_text[-1] == "=" * 40:
            output_text.pop()

        return "\n".join(output_text).strip()


    # --- Education/Certifications Format (NOW GROUPED AND SORTED) ---
    elif destination == "Education":
        
        # Robust column fetching for Education/Certifications
        org_col = 'Authority' if 'Authority' in df.columns else 'Issuing Organization'
        if org_col not in df.columns and 'School Name' in df.columns:
            org_col = 'School Name'
        
        # Determine which column holds the name/title of the education/cert
        name_col = 'Name' if 'Name' in df.columns else 'Degree Name'
        if name_col not in df.columns:
             name_col = 'School Name' 

        start_col = 'Started On' if 'Started On' in df.columns else 'Start Date'
        finish_col = 'Finished On' if 'Finished On' in df.columns else 'End Date'

        output_groups = {}
        
        # 1. Prepare for Sorting/Grouping
        df['Category'], df['Priority'] = zip(*df.apply(classify_education_entry, axis=1))
        df['End_Date_Obj'] = df[finish_col].apply(parse_date_for_sort)
        df['Start_Date_Obj'] = df[start_col].apply(parse_date_for_sort)
        
        # 2. Sort: Priority (Academic > Cert) Descending, then Date (Most Recent) Descending
        df = df.sort_values(by=['Priority', 'End_Date_Obj', 'Start_Date_Obj'], ascending=[False, False, False])

        # 3. Format and Group
        for _, row in df.iterrows():
            final_name_raw = str(row.get('Degree Name', '')) if 'Degree Name' in df.columns and str(row.get('Degree Name', '')).strip() else str(row.get(name_col, ''))
            final_name = final_name_raw.strip()
            
            if not final_name or final_name.lower() == 'nan': continue

            org = str(row.get(org_col, ''))
            if org.lower() == 'nan': org = ''
            
            start_date = str(row.get(start_col, '')).split('T')[0].strip()
            finish_date = str(row.get(finish_col, '')).split('T')[0].strip()

            date_info = []
            if start_date and start_date.lower() != 'nan': date_info.append(start_date)
            if finish_date and finish_date.lower() != 'nan': date_info.append(finish_date)

            date_str = " - ".join(date_info)
            if not date_str: date_str = "Date Unknown"

            category = row['Category']
            formatted_line = f"**{final_name}** | {org.strip()} ({date_str})"
            
            output_groups.setdefault(category, []).append(formatted_line)
        
        education_output = ["--- NEW EDUCATION & CERTIFICATIONS ---"]
        
        # Custom Order for display (Highest Priority first)
        category_order = [
            "Academic Degrees (PhD/Doctorate)", "Academic Degrees (Master's)", "Academic Degrees (Bachelor's & High School)",
            "Professional Certifications (Data Science/ML)", "Professional Certifications (Business/Leadership)",
            "Professional Certifications (Engineering)", "Professional Certifications (Teaching/Language)", 
            "Professional Certifications (Other)"
        ]

        for cat in category_order:
            if cat in output_groups:
                education_output.append(f"\n### {cat.upper()}:")
                education_output.extend([f"- {line}" for line in output_groups[cat]])
                del output_groups[cat]
        
        # Add any remaining (unlikely with comprehensive categories)
        for cat, lines in output_groups.items():
            education_output.append(f"\n### {cat.upper()}:")
            education_output.extend([f"- {line}" for line in lines])

        return "\n".join(education_output).strip()


    # --- Skills Format (NOW GROUPED HEURISTICALLY) ---
    elif destination == "Skills":
        skill_col = 'Skill Name' if 'Skill Name' in df.columns else 'Name'
        
        if skill_col not in df.columns:
             return "Error: Skills CSV missing expected column ('Skill Name' or 'Name')."
             
        # Extract, convert to string, strip whitespace, and filter out empty or 'nan' strings
        skills_raw = df[skill_col].astype(str).loc[
            df[skill_col].astype(str).str.strip().str.lower() != 'nan'
        ].loc[
            df[skill_col].astype(str).str.strip() != ''
        ].tolist()
        
        # Heuristic Groups
        groups = {
            "Spoken Languages": [],
            "Programming & Data": [],
            "Cloud & IT Systems": [],
            "Tools & Software": [],
            "Soft Skills & Management": [],
            "Other Technical Skills": []
        }
        
        # Keywords for grouping (can be expanded)
        programming_kws = ['python', 'r', 'java', 'javascript', 'c++', 'c#', 'sql', 'html', 'css', 'go', 'php', 'swift', 'data science', 'machine learning', 'mlops', 'tableau', 'power bi']
        languages_kws = ['english', 'spanish', 'french', 'german', 'japanese', 'dutch', 'korean', 'mandarin']
        cloud_kws = ['aws', 'azure', 'gcp', 'cloud', 'virtualization', 'networking']
        tools_kws = ['jira', 'confluence', 'git', 'excel', 'word', 'adobe', 'office 365', 'cad']
        soft_kws = ['leadership', 'management', 'communication', 'teamwork', 'problem-solving', 'analytical', 'negotiation', 'training', 'teaching', 'first-aider', 'firefighter']

        for skill in skills_raw:
            s_lower = skill.lower()
            
            if any(kw in s_lower for kw in languages_kws):
                groups["Spoken Languages"].append(skill)
            elif any(kw in s_lower for kw in programming_kws):
                groups["Programming & Data"].append(skill)
            elif any(kw in s_lower for kw in cloud_kws):
                groups["Cloud & IT Systems"].append(skill)
            elif any(kw in s_lower for kw in tools_kws):
                groups["Tools & Software"].append(skill)
            elif any(kw in s_lower for kw in soft_kws):
                groups["Soft Skills & Management"].append(skill)
            else:
                groups["Other Technical Skills"].append(skill)
                
        output_text = ["--- NEW SKILLS ---"]
        for group, skills in groups.items():
            if skills:
                output_text.append(f"\n### {group.upper()}:")
                # Use a comma-separated list for skills within a group
                output_text.append(", ".join(sorted(list(set(skills))))) 
                
        return "\n".join(output_text).strip()
            
    return f"Error: Unknown destination **{destination}** in formatter."


def handle_final_csv_processing(csv_file_path, destination, action, *current_boxes):
    """Processes the CSV after user selects the destination and action."""
    
    # Map destination string to the index of the profile boxes (0-5)
    DESTINATION_MAP = {
        "Contact": 0, "Summary": 1, "Experience": 2, 
        "Education": 3, "Skills": 4, "Projects": 5
    }
    
    # Handle 'Ignore' selection
    if destination == "Ignore":
        status = "⚠️ CSV processing ignored. File data discarded."
        return list(current_boxes) + [status] + [None] + [gr.update(interactive=False)] * 3
        
    
    # 1. Format the new data based on user's selected destination
    formatted_text = _format_csv_data(csv_file_path, destination)
    
    if isinstance(formatted_text, str) and formatted_text.startswith("Error"):
        status = f"Formatting Error for **{destination}**: {formatted_text}"
        return list(current_boxes) + [status] + [None] + [gr.update(interactive=False)] * 3

    # 2. Get the index and current content
    idx = DESTINATION_MAP.get(destination)
    if idx is None:
        status = f"Error: Invalid destination selected: **{destination}**."
        return list(current_boxes) + [status] + [None] + [gr.update(interactive=False)] * 3
        
    current_content = list(current_boxes)
    updates = [gr.update(value=c) for c in current_content] 
    
    new_data_string = formatted_text.strip()
    
    if action == "Overwrite (Replace)":
        updates[idx] = gr.update(value=new_data_string)
        status = f"✅ Section **{destination}** successfully **OVERWRITTEN**."
        
    elif action == "Append (Add)":
        # Professional Append: Add a clear separator and the new data
        current_data = current_content[idx].strip()
        
        # Use a standard separator that is stripped by the MD loader
        if current_data:
            # Separator used for Experience and Projects is different
            separator_char = "=" if destination in ["Education", "Projects"] else "-"
            separator = "\n\n" + (separator_char * 80) + "\n\n"
            
            # Remove the "--- NEW SECTION ---" header from the appended data if current data exists
            if current_data and new_data_string.startswith("--- NEW "):
                 new_data_string = new_data_string.split('\n', 1)[-1].strip()

            final_text = current_data + separator + new_data_string
        else:
            final_text = new_data_string
            
        updates[idx] = gr.update(value=final_text)
        status = f"✅ Section **{destination}** successfully **APPENDED** with new data."
        
    # 3. Disable controls and reset state
    return updates + [status] + [None] + [gr.update(interactive=False)] * 3


def store_structured_profile(contact, summary, experience, education, skills, projects):
    """Saves all six profile fields to their respective state components."""
    if not experience.strip() and not skills.strip():
        return ("⚠️ Saved: Profile is nearly empty. Please fill in Experience or Skills.", 
                contact, summary, experience, education, skills, projects,
                contact, summary, experience, education, skills, projects)
        
    return (
        "✅ Saved: Master Profile successfully updated from manual input. Ready for generation.", 
        contact, summary, experience, education, skills, projects,  
        contact, summary, experience, education, skills, projects  
    )

def get_profile_data(contact, summary, experience, education, skills, projects):
    """Combines all state values into a single master resume string for the LLM."""
    if not experience and not skills:
        return "Error: Profile is empty. Please fill in the Master Profile tab and click 'Save'."
        
    # Headers used here must exactly match those used in parse_md_file for loading to work
    # Ensured extra newlines for clean separation in the prompt
    master_resume = (
        f"--- CANDIDATE MASTER PROFILE START (Comprehensive Data) ---\n"
        f"CONTACT:\n{contact}\n\n"
        f"SUMMARY:\n{summary}\n\n"
        f"EXPERIENCE (Detailed Career History & Achievements):\n{experience}\n\n"
        f"PROJECTS (Key personal/work projects, Awards, Publications):\n{projects}\n\n" 
        f"EDUCATION (Degrees and Certifications):\n{education}\n\n"
        f"SKILLS (All Technical and Soft Skills):\n{skills}\n"
        f"--- CANDIDATE MASTER PROFILE END ---\n"
    )
    return master_resume

# --- ARCHIVAL FUNCTION ---
def archive_master_profile_wrapper(contact, summary, experience, education, skills, projects, format_choice):
    """Combines the full master profile data and formats it for download."""
    master_text = get_profile_data(contact, summary, experience, education, skills, projects)
    
    if master_text.startswith("Error"): 
        return None, master_text
        
    file_path = generate_formatted_output(master_text, "MASTER_PROFILE_ARCHIVE", "Raw Data Dump", format_choice)
    
    if file_path:
        return file_path, f"✅ Master Profile Archive file generated successfully in {format_choice} format. Upload this file later to quickly reload your profile."
    else:
        return None, f"⚠️ Failed to generate archive file. Check console for errors."

# --------------------------------------------------------------------------
# --- LLM WRAPPER & FORMATTING FUNCTIONS ---
# --------------------------------------------------------------------------

# Assuming your function definition looks something like this:
def get_match_score_wrapper(target_position, job_description, language, contact, summary, experience, education, skills, projects):
    
    # 1. ENRICH SKILLS
    # This step is VITAL to include tools from certifications (like Jupyter, Matplotlib)
    updated_skills = enrich_skills_from_certifications(skills, education)
    
    # 2. CREATE MASTER PROFILE with the updated skills list
    master_profile = get_profile_data(contact, summary, experience, education, updated_skills, projects)
    
    if master_profile.startswith("Error"):
        return master_profile

    # 3. LLM SCORING PROMPT (with improved formatting rules)
    score_prompt = f"""
    You are an expert ATS (Applicant Tracking System) Analyst. Compare the following Candidate Master Profile against the Job Description.

    [...] (The rest of your scoring prompt) [...]

    CRITICAL OUTPUT FORMATTING RULE:
    1.  Start the output with the score on a new line: "SCORE: **[Calculated Score]**"
    2.  Use Markdown headings (##) for the three main sections.
    3.  Use bullet points (-) for every item under the "Missing Keywords" section.
    4.  Ensure every section is separated by a blank line for readability.
    """
    
    final_output = llm_chat(score_prompt, max_tokens=1024)
    return final_output

def polish_resume(target_position, job_description, polish_prompt, resume_style, language, contact, summary, experience, education, skills, projects):
    """Generates a polished and complete resume document with enriched skills."""
    
    if not target_position or not job_description:
        return "Error: Please provide a Target Position and Job Description.", gr.update(interactive=False), ""
    
    # CRITICAL NEW STEP: Use the helper function to enrich the skills list
    # The helper function uses the LLM to read 'education' and merge new skills into 'skills'.
    updated_skills = enrich_skills_from_certifications(skills, education)
    
    # Now, call get_profile_data using the newly ENRICHED skill list
    master_resume = get_profile_data(contact, summary, experience, education, updated_skills, projects)
    
    if master_resume.startswith("Error"):
        # The third return value should be "" or the error if docx generation is expected, 
        # but matching the previous style:
        return master_resume, gr.update(interactive=False), master_resume

    # UNIFIED PROMPT to generate the complete resume structure
    prompt_use = f"""
    As an expert Resume Writer, your task is to generate a **complete, professional resume document** tailored for the **{target_position}** role.

    The resume must be one page (if possible) and must include **all** standard, fully formatted resume sections:
    1.  **CONTACT** (Use the contact details from the Master Profile).
    2.  **SUMMARY/PROFILE** (Targeted to the job description).
    3.  **EXPERIENCE** (Rewritten to highlight achievements and skills most relevant to the job description).
    4.  **EDUCATION & CERTIFICATIONS** (Filter and prioritize the most relevant degrees and certifications. **Exclude similar or outdated certifications, focusing only on those directly supporting the role.**).
    5.  **SKILLS** (Grouped logically and prioritized based on the job description).

    **Style**: The format must strictly follow the **{resume_style}** standard.
    **Specific Instruction**: {polish_prompt if polish_prompt else 'Focus on action verbs and quantifiable results.'}
    **CRITICAL LANGUAGE RULE**: The entire resume must be in **{language}**.

    Candidate Master Profile (All Data): '''{master_resume}'''
    Job Description (for Targeting): '''{job_description}'''

    CRITICAL OUTPUT RULE: Return ONLY the complete, fully formatted resume text. Do not include any commentary or preamble.
    """

    # Use the larger token limit (2048) for the full generation
    final_output = llm_chat(prompt_use, max_tokens=2048) 

    if final_output.startswith("Error"):
        return final_output, gr.update(interactive=False), final_output
    
    # Note: The original snippet was missing the docx saving logic, 
    # ensure you re-add the file generation here if you haven't already.
    # return final_output, gr.update(interactive=True), final_output # Original line for return
    
    # Assuming the final file generation/return logic is correct in your 1311-line file:
    return final_output, gr.update(interactive=True), final_output

def generate_cover_letter(company_name, target_position, job_description, cover_letter_style, language, contact, summary, experience, education, skills, projects):
    """Generates a customized cover letter."""
    
    # --- 1. Initial Validation and Data Retrieval ---
    if not company_name or not target_position or not job_description:
        return "Error: Please provide Company Name, Position Name, and Job Description.", gr.update(interactive=False), ""
        
    master_resume = get_profile_data(contact, summary, experience, education, skills, projects)
    if master_resume.startswith("Error"): 
        return master_resume, gr.update(interactive=False), "" # Return empty file path on error
        
    # CRITICAL FIX: Extract the candidate name immediately after getting the contact data
    candidate_name = extract_name(contact) 
    
    contact_info = contact if contact.strip() else "Your Name and Full Contact Details (Please update the Contact Info box!)"
    
    # --- 2. LLM Prompt Construction ---
    prompt = f"""
    Generate a professional, customized cover letter for the **{target_position}** role at **{company_name}** in a professional tone.
    
    **Letter Style**: The letter MUST strictly follow the **{cover_letter_style}** standards.
    **Candidate Info**: Use the contact info provided: {contact_info}.
    **Content**: The letter must directly reference and integrate the candidate's specific qualifications and experience (provided in the Master Profile) to show a strong fit for the Job Description.
    
    CRITICAL LANGUAGE RULE: The entire letter must be written in **{language}** (e.g., Japanese, English, French).
    
    Master Profile: '''{master_resume}'''
    Job Description: '''{job_description}'''
    Candidate Contact Info: '''{contact_info}'''
    
    CRITICAL OUTPUT RULE: Return only the complete cover letter text, formatted with appropriate paragraphs and line breaks, ready to be sent. Do not include any commentary or preamble.
    """
    
    # --- 3. LLM Call ---
    final_output = llm_chat(prompt, max_tokens=2048)

    # --- 4. Final Processing and File Generation ---
    if final_output.startswith("Error"): 
        return final_output, gr.update(interactive=False), ""
        
    # Use the extracted candidate_name for safe file naming
    output_filename = f"{candidate_name.replace(' ', '_')}_Cover_Letter_for_{target_position.replace(' ', '_')}.docx"
    docx_file_path = save_docx(final_output, output_filename)

    # If save_docx returns an error string, use that, otherwise success
    if docx_file_path.startswith("Error"):
        return final_output + "\n\n" + docx_file_path, gr.update(interactive=False), ""

    return final_output, gr.update(interactive=True), docx_file_path

def generate_linkedin_strategy(target_position, language, contact, summary, experience, education, skills, projects):
    """Generates universal LinkedIn profile advice focused on multi-industry appeal."""

    # You would still compile your master profile data here
    master_profile = get_profile_data(contact, summary, experience, education, skills, projects)
    if master_profile.startswith("Error"):
        return master_profile, gr.update(interactive=False), ""
    
    # CRITICAL CHANGE: The prompt removes all references to a specific Job Description.
    prompt_use = f"""
    You are an expert LinkedIn Career Strategist. Your task is to analyze the candidate's Master Profile and generate a complete, strategic LinkedIn profile overhaul for the **{target_position}** profession.

    The primary goal is to create a **universal profile** that is effective for searching and appealing to recruiters across **multiple industries** (e.g., Healthcare, Tech, Finance, E-commerce).

    **CRITICAL OUTPUT FORMATTING INSTRUCTIONS:**
    1.  **Use Markdown Headings (##) and bolding (**) for all section titles.**
    2.  Ensure liberal use of **paragraph breaks** in the About section for readability.
    3.  Use **Markdown bullet points (- )** for the reframed experiences section.

    Provide the output in three distinct sections:

    ## 1. Headline Strategy (Max 150 characters)
    Generate the ideal, keyword-rich professional headline. Focus on impact and core methodology, not industry. (e.g., 'Data Scientist | Predictive Modeling & AI/ML | Driving Business Value')

    ## 2. About Section (Full Draft)
    Write a compelling 'About' section (4-5 paragraphs).
    * **CRITICAL:** Integrate transferable skills and **multi-industry keywords** to maximize searchability (e.g., applying ML to 'financial risk,' 'supply chain optimization,' or 'customer personalization').
    * Use strong, quantifiable achievement statements from the candidate's Experience/Projects.

    ## 3. Experience & Projects Reframing (Bulleted List)
    For the candidate's top 3 most impressive experiences/projects, provide a new, universal, and quantifiable **one-paragraph summary for each.** **Present these three reframed summaries as a Markdown bulleted list.**

    CRITICAL LANGUAGE RULE: The entire output must be written in **{language}**.

    Candidate Master Profile (All Data): '''{master_profile}'''

    CRITICAL OUTPUT RULE: Return ONLY the three requested sections of advice, formatted strictly using Markdown headings and lists. Do not include any commentary or preamble.
    """

    final_output = llm_chat(prompt_use, max_tokens=2048) 

    if final_output.startswith("Error"):
        return final_output, gr.update(interactive=False), ""

    # Since this is advice (text), you likely don't need the save_docx call here, 
    # but you would return the text output and an empty file component.
    return final_output, gr.update(interactive=False), ""

def analyze_resume_match(master_resume, job_description, language):
    """Generates the ATS match score and analysis."""
    prompt = f"""
    You are an expert ATS (Applicant Tracking System) Analyst. Analyze the overlap between the candidate's Master Profile and the Job Description.
    
    Master Profile: '''{master_resume}'''
    Job Description: '''{job_description}'''
    
    CRITICAL LANGUAGE RULE: The entire analysis (including headers, lists, and summary) must be written in **{language}**.
    
    1. **Score:** Provide a single match score out of 100 based purely on keyword and skill overlap. Output this as the first line: `SCORE: [0-100]%`
    
    2. **Missing Keywords:** List the top 5 most critical keywords/skills from the Job Description that are missing or underrepresented in the profile.
    
    3. **Actionable Summary:** Provide a short, actionable summary (3-4 sentences) on how the candidate can increase their score for this specific job.
    
    Return the SCORE line first, followed by the rest of the analysis in a structured, easy-to-read format.
    """
    if not job_description: return "Error: Please paste the Job Description to get a match score."
    return llm_chat(prompt)

def analyze_interview_feedback(history, target_position, job_description, language):
    """Analyzes the interview history and provides coaching feedback."""
    if len(history) < 2: return "Error: Please complete at least one question-answer cycle before requesting feedback."
    
    conversation_lines = []
    for q_a_pair in history:
        if q_a_pair[1]: 
             conversation_lines.append(f"Q: {q_a_pair[1]}") 
        if q_a_pair[0]:
            conversation_lines.append(f"A: {q_a_pair[0]}")
            
    if not any(line.startswith("A:") for line in conversation_lines):
         return "Error: Please complete at least one question-answer cycle before requesting feedback."
        
    conversation = "\n".join(conversation_lines)
    
    prompt = f"""
    You are an AI Interview Coach. Analyze the following practice interview transcript for the '{target_position}' role.
    
    Conversation Transcript:
    ---
    {conversation}
    ---
    
    CRITICAL LANGUAGE RULE: The entire feedback must be written in **{language}**.
    
    1. **Overall Score:** Provide a score out of 10 based on clarity, relevance to the JD, and impact. Output this as the first line: `SCORE: [0-10]/10`
    
    2. **Weakest Response:** Identify the single weakest answer or biggest missed opportunity in the candidate's responses. Quote the specific question and the candidate's answer.
    
    3. **Rephrasing Advice:** Provide a concise, professional rephrased version of that weakest response (using the STAR method if applicable), explaining *why* the new phrasing is stronger.
    
    Return the SCORE line first, followed by the rest of the analysis.
    """
    return llm_chat(prompt)

def interview_chat(message, history, target_position, job_description, interview_type, language, profile_data):
    """The core logic for the Gradio ChatInterface, powered by LLM."""
    
    if profile_data is None or profile_data.startswith("Error"):
        if len(history) == 0:
            return "Please click **'1. Start/Reset Interview (Loads Profile)'** before beginning to ensure the AI Recruiter has your profile context."
        
    system_instruction = f"""
    You are an expert, professional recruiter conducting a practice interview for the position of '{target_position}'.
    The candidate's full profile is: '''{profile_data}'''
    The job description is: '''{job_description}'''
    The interview type is: {interview_type}.
    
    CRITICAL LANGUAGE RULE: All your questions and responses must be in **{language}**.
    
    Your role is to: 1. Base questions on the provided candidate profile and job description. 2. Maintain a professional, consistent persona. 3. Be concise and move the interview forward after each response. 4. Do not offer feedback during the interview.
    """
    
    llm_history = [{"role": "system", "content": [{"type": "text", "text": system_instruction}]}]
    
    for user_msg, assistant_msg in history:
        if assistant_msg:
             llm_history.append({"role": "assistant", "content": [{"type": "text", "text": assistant_msg}]})
        if user_msg:
             llm_history.append({"role": "user", "content": [{"type": "text", "text": user_msg}]})

    current_message = {"role": "user", "content": [{"type": "text", "text": message}]}
    
    try:
        generated_response = model.chat(messages=llm_history + [current_message])
        return generated_response['choices'][0]['message']['content']
    except Exception as e:
        return f"Recruiter Error: {e}"
        
# --------------------------------------------------------------------------
# --- FORMATTING FUNCTION (FIXED FOR STABILITY AND NAMING) ---
# --------------------------------------------------------------------------
def generate_formatted_output(output_text, output_type, resume_style, format_choice):
    """Generates and saves the output text in the chosen format (DOCX, MD) with robust formatting."""
    if not output_text or output_text.startswith("Error"): 
        return None
    
    # Global list tracking temp files
    global temp_files_to_clean
    base_filename = f"output_{output_type.lower().replace(' ', '_')}_{os.getpid()}"
    
    if "markdown" in format_choice.lower() or "md" in format_choice.lower():
        # --- Stable Markdown/Plain Text Output ---
        ext = "md"
        final_file_path = f"{base_filename}.{ext}"
        try:
            # Simple file write for maximum stability
            with open(final_file_path, 'w', encoding='utf-8') as f:
                # LLM output is already in markdown-style, just write it
                f.write(output_text)
            temp_files_to_clean.append(final_file_path)
            return final_file_path
        except Exception as e:
            print(f"Error saving Markdown file: {e}")
            return None
            
    elif "docx" in format_choice.lower():
        # --- Unstable DOCX Output (Kept for completeness, using the simplified structure) ---
        ext = "docx"
        temp_file_path = f"{base_filename}.{ext}" 
        
        document = DocxDocument()
        
        # Set basic style for the whole document
        style = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Title/Header
        title_p = document.add_paragraph()
        title_p.add_run(f"{output_type.replace('_', ' ').title()} ({resume_style})").bold = True
        title_p.style.font.size = Pt(14)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph() # Add a blank line

        # Process text line-by-line with aggressive simplification
        for line in output_text.split('\n'):
            cleaned_line = line.strip()
            if not cleaned_line: 
                document.add_paragraph()
                continue
            
            # 1. Handle Headings (e.g., ### DATA SCIENCE or --- NEW SECTION ---)
            if cleaned_line.startswith('###') or cleaned_line.startswith('--- NEW') or ('--' * 40) in cleaned_line or ('==' * 40) in cleaned_line:
                header_text = cleaned_line.strip('#').strip().strip('-').strip().strip('=').strip()
                # Use a standard paragraph with bolding instead of a DOCX heading style for maximum stability
                p = document.add_paragraph()
                p.add_run(header_text).bold = True
                p.style.font.size = Pt(12)
                p.space_after = Pt(6)
            
            # 2. Handle All Other Text (Including lists)
            else:
                p = document.add_paragraph()
                # AGGRESSIVE SIMPLIFICATION: We use regex to split ONLY on **bold** markers
                parts = re.split(r'(\*\*.*?\*\*)', cleaned_line)
                
                for part in parts:
                    # Handle **bolded text**
                    if part.startswith('**') and part.endswith('**'):
                        # Create a run that is ONLY bold
                        p.add_run(part.strip('*')).bold = True
                    else:
                        # Create a run that is normal text
                        p.add_run(part)

        # Save the generated content as a DOCX file
        document.save(temp_file_path)
        temp_files_to_clean.append(temp_file_path)
        return temp_file_path

    # Fallback/Error if format is unexpected
    return None

# --------------------------------------------------------------------------
# --- GRADIO INTERFACE SETUP ---
# --------------------------------------------------------------------------

# New choices for stable output, including the preferred Markdown
STABLE_FORMAT_CHOICES = ["Markdown (.md)", "DOCX (Unstable)"]

with gr.Blocks(title="AI Career Coach Pro") as demo:
    gr.Markdown("## 💼 AI Career Coach Pro: Complete Toolkit for Job Seekers")
    
    # Global Language Selector
    LANGUAGES = ["English", "Japanese", "Spanish", "French", "German"]
    with gr.Row():
        language_selector = gr.Radio(
            label="🎯 Select Output Language for ALL LLM Tasks (Except Parsing)", 
            choices=LANGUAGES, 
            value="English",
            scale=3
        )
    
    # State components (Internal storage for profile data)
    MASTER_CONTACT_STATE = gr.State("")
    MASTER_SUMMARY_STATE = gr.State("")
    MASTER_EXPERIENCE_STATE = gr.State("")
    MASTER_EDUCATION_STATE = gr.State("")
    MASTER_SKILLS_STATE = gr.State("")
    MASTER_PROJECTS_STATE = gr.State("") 
    LAST_GENERATED_TEXT = gr.State("") 
    INTERVIEW_PROFILE_DATA = gr.State(get_profile_data("", "", "", "", "", ""))
    
    # --- NEW STATE FOR CSV PROCESSING ---
    CSV_FILE_TO_PROCESS = gr.State(None)

    RESUME_STYLES = ["ATS Optimized", "Chronological", "Functional/Skills-Based", "Hybrid"]
    CL_STYLES = ["ATS Optimized", "Traditional Formal", "Modern Brief"]
    state_inputs = [MASTER_CONTACT_STATE, MASTER_SUMMARY_STATE, MASTER_EXPERIENCE_STATE, MASTER_EDUCATION_STATE, MASTER_SKILLS_STATE, MASTER_PROJECTS_STATE]

    # --- Master Profile Tab ---
    with gr.Tab("👤 Master Profile (Source of Truth)"):
        gr.Markdown("### 1. Upload & Populate Fields")
        gr.Markdown("Upload **PDF/TXT** resume for AI parsing, **LinkedIn CSV** files, or a saved **Markdown (.md) Archive** to load your profile.")
        
        master_file_input = gr.File(label="Upload Resume (PDF/TXT), LinkedIn Data (CSV), or Archive (.md)", file_types=[".pdf", ".txt", ".csv", ".md"], type="filepath")
        status_output_profile = gr.Textbox(label="Status Log", lines=2, interactive=False)
        
        gr.Markdown("### 2. Ambiguous CSV Processing (Required if Status Log requests action)")
        with gr.Row():
            destination_section_radio = gr.Radio(
                label="Destination Section", 
                choices=["Experience", "Projects", "Education", "Skills", "Ignore"], 
                value="Ignore",
                interactive=False,
                scale=2
            )
            action_radio = gr.Radio(
                label="Action", 
                choices=["Overwrite (Replace)", "Append (Add)"], 
                value="Append (Add)", # Defaulting to Append is safer for users
                interactive=False,
                scale=1
            )
            finalize_csv_button = gr.Button("Process Selected CSV", variant="primary", interactive=False, scale=2)
        
        gr.Markdown("### 3. Review and Edit Profile Sections")
        
        with gr.Row():
            contact_box = gr.Textbox(label="Contact & Personal Info", lines=4, interactive=True)
            summary_box = gr.Textbox(label="Summary/Objective", lines=4, interactive=True)
        
        experience_box = gr.Textbox(label="Core Work Experience & Achievements (Ordered by Date)", lines=15, interactive=True)
        
        with gr.Row():
            education_box = gr.Textbox(label="Education & Certifications (Grouped by Diploma/Cert)", lines=10, interactive=True) 
            skills_box = gr.Textbox(label="Technical Skills (Grouped by Category)", lines=10, interactive=True) 
            
        projects_box = gr.Textbox(label="Projects/Awards/Publications", lines=5, interactive=True) 

        textbox_outputs = [contact_box, summary_box, experience_box, education_box, skills_box, projects_box]

        save_button = gr.Button("💾 Save All Manual Changes to Profile", variant="primary")
        
        gr.Markdown("---")
        
        # --- ARCHIVAL SECTION ---
        gr.Markdown("### 4. Archive Full Profile (Your Future-Proof File)")
        with gr.Row():
            # Renaming the choice to reflect DOCX's ability to be converted to PDF
            download_format_archive = gr.Radio(label="Archive Format", choices=["Markdown (.md)", "DOCX (for PDF conversion)"], value="Markdown (.md)", scale=1)
            archive_button = gr.Button("⬇️ Download Full Profile Archive", variant="secondary", scale=2)
        pdf_file_output_archive = gr.File(label="Download Master Archive File", scale=3)
        
        archive_button.click(
            fn=archive_master_profile_wrapper,
            inputs=state_inputs + [download_format_archive],
            outputs=[pdf_file_output_archive, status_output_profile]
        )
        # --- END ARCHIVAL SECTION ---

        # WIRING THE NEW CSV/MD LOGIC
        # Output order: 6 textboxes, status, file_state, radio_destination, radio_action, button
        upload_outputs = textbox_outputs + [status_output_profile] + [CSV_FILE_TO_PROCESS, destination_section_radio, action_radio, finalize_csv_button]
        
        master_file_input.upload(
            fn=parse_and_distribute_file, 
            inputs=[master_file_input], 
            outputs=upload_outputs
        )
        
        # Inputs for finalization: file_state, radio_destination, radio_action, PLUS the current 6 textboxes
        finalize_csv_button.click(
            fn=handle_final_csv_processing,
            inputs=[CSV_FILE_TO_PROCESS, destination_section_radio, action_radio] + textbox_outputs,
            outputs=upload_outputs
        )
        
        # STANDARD SAVE BUTTON
        save_button.click(
            fn=store_structured_profile, 
            inputs=textbox_outputs, 
            outputs=[status_output_profile] + textbox_outputs + state_inputs
        )

    # --- Resume Polisher Tab ---
    with gr.Tab("📝 Resume Polisher"):
        gr.Markdown("### 1. Inputs")
        target_pos_polisher = gr.Textbox(label="Target Position Name", placeholder="e.g., Senior Data Scientist")
        jd_polisher = gr.Textbox(label="Job Description (Required)", placeholder="Paste the full job description here...", lines=10)
        polish_inst = gr.Textbox(label="Specific Polish Instructions (Optional)", placeholder="e.g., 'Use more action verbs' or 'limit to one page'", lines=3)
        style_radio_polisher = gr.Radio(label="Output Style", choices=RESUME_STYLES, value="ATS Optimized")

        gr.Markdown("### 2. Match Score Analysis")
        score_button = gr.Button("Get Resume Match Score & Analysis", variant="secondary")
        match_score_output = gr.Textbox(label="ATS Match Analysis", lines=10, interactive=False)
        
        score_button.click(
            fn=get_match_score_wrapper,
            inputs=[target_pos_polisher, jd_polisher, language_selector] + state_inputs,
            outputs=match_score_output
        )

        gr.Markdown("### 3. Generate and Download")
        with gr.Row():
            polish_button = gr.Button("Generate Polished Text", variant="secondary", scale=1)
            download_format_polisher = gr.Radio(label="Choose Download Format", choices=["Markdown (.md)", "DOCX (for PDF conversion)"], value="Markdown (.md)", scale=1)
            download_button_polisher = gr.Button("Download Formatted File", variant="primary", interactive=False, scale=1)
        
        polished_text_output = gr.Textbox(label="Text Output (For Review)", lines=15, interactive=False)
        pdf_file_output_polisher = gr.File(label="Download Formatted File")
        
        polish_button.click(
            fn=polish_resume, 
            inputs=[target_pos_polisher, jd_polisher, polish_inst, style_radio_polisher, language_selector] + state_inputs, 
            outputs=[polished_text_output, download_button_polisher, LAST_GENERATED_TEXT]
        )

        download_button_polisher.click(
            fn=generate_formatted_output, 
            inputs=[LAST_GENERATED_TEXT, gr.State("RESUME_SNIPPET"), style_radio_polisher, download_format_polisher], 
            outputs=pdf_file_output_polisher
        )

    # --- Cover Letter Generator Tab ---
    with gr.Tab("✉️ Cover Letter Generator"):
        gr.Markdown("### 1. Inputs")
        company_name_cl = gr.Textbox(label="Company Name", placeholder="e.g., IBM")
        position_name_cl = gr.Textbox(label="Position Name", placeholder="e.g., AI Developer")
        jd_cl = gr.Textbox(label="Job Description (Required)", placeholder="Paste the full job description here...", lines=10)
        style_radio_cl = gr.Radio(label="Output Style", choices=CL_STYLES, value="Traditional Formal")
        
        gr.Markdown("### 2. Generate and Download")
        with gr.Row():
            cl_button = gr.Button("Generate Cover Letter Text", variant="secondary", scale=1)
            download_format_cl = gr.Radio(label="Choose Download Format", choices=["Markdown (.md)", "DOCX (for PDF conversion)"], value="Markdown (.md)", scale=1)
            download_button_cl = gr.Button("Download Formatted File", variant="primary", interactive=False, scale=1)
        
        cl_output = gr.Textbox(label="Text Output (For Review)", lines=15, interactive=False)
        pdf_file_output_cl = gr.File(label="Download Formatted File")

        cl_button.click(
            fn=generate_cover_letter, 
            inputs=[company_name_cl, position_name_cl, jd_cl, style_radio_cl, language_selector] + state_inputs, 
            outputs=[cl_output, download_button_cl, LAST_GENERATED_TEXT]
        )
        
        download_button_cl.click(
            fn=generate_formatted_output, 
            inputs=[LAST_GENERATED_TEXT, gr.State("COVER_LETTER"), style_radio_cl, download_format_cl], 
            outputs=pdf_file_output_cl
        )

    # --- LinkedIn & Networking Advisor Tab ---
    with gr.Tab("🌐 LinkedIn & Networking Advisor"):
        gr.Markdown("### Strategic LinkedIn Profile & Career Advice")
        target_pos_advice = gr.Textbox(label="Target Position Applied For", placeholder="e.g., Machine Learning Engineer")        
        
        advice_button = gr.Button("Generate LinkedIn Strategy & Advice", variant="primary")
        advice_output = gr.Textbox(label="AI Generated Output", lines=30, interactive=False)
        
        advice_button.click(
            fn=generate_linkedin_strategy, 
            inputs=[target_pos_advice, language_selector] + state_inputs, 
            outputs=advice_output
        )

    # --- Practice Interview Tab ---
    with gr.Tab("🗣️ Practice Interview"):
        gr.Markdown("### 1. Setup")
        
        with gr.Row():
            target_pos_interview = gr.Textbox(label="Target Position Name", placeholder="e.g., Senior Data Scientist", scale=1)
            interview_type_radio = gr.Radio(
                label="Interview Type", 
                choices=["Behavioral", "Technical/Situational", "General Screening"],
                value="Behavioral",
                scale=1
            )
            
        jd_interview = gr.Textbox(label="Job Description (Required for Context)", placeholder="Paste the full job description here...", lines=5)
        
        interview_status = gr.Textbox(label="Setup Status", value="Profile not loaded. Click 'Start/Reset' to begin.", interactive=False)
        prep_button = gr.Button("1. Start/Reset Interview (Loads Profile)", variant="secondary")

        gr.Markdown("### 2. Live Interview")

        interview_chat_interface = gr.ChatInterface(
            fn=lambda message, history: interview_chat(
                message, 
                history, 
                target_pos_interview.value, 
                jd_interview.value, 
                interview_type_radio.value,
                language_selector.value, 
                INTERVIEW_PROFILE_DATA.value
            ),
            chatbot=gr.Chatbot(label="AI Recruiter", height=400, type='messages'),
            textbox=gr.Textbox(placeholder="Your response...", lines=4),
            submit_btn="Send Response",
        )
        
        gr.Markdown("### 3. Coaching Feedback")
        score_interview_button = gr.Button("Get Interview Score & Coaching Feedback", variant="primary", interactive=False)
        interview_feedback_output = gr.Textbox(label="Interview Feedback & Rephrasing Advice", lines=10, interactive=False)

        prep_button.click(
            fn=get_profile_data,
            inputs=state_inputs,
            outputs=[INTERVIEW_PROFILE_DATA]
        ).then(
            fn=lambda profile_data, target_pos, int_type, jd, lang: [
                gr.update(interactive=True), 
                gr.update(value=f"✅ Profile loaded for **{target_pos}** (Type: {int_type}). Chat cleared. Language: {lang}. Ready for the first question."), 
                [], 
                {'role': 'assistant', 'content': f"Hello, I'm the AI Recruiter. I've loaded your profile for the **{target_pos}** role. As we conduct this interview in **{lang}**, let's begin the {int_type} interview. Can you start by telling me a little about your background, focusing on {jd.split('.')[0]}?"}
            ], 
            inputs=[INTERVIEW_PROFILE_DATA, target_pos_interview, interview_type_radio, jd_interview, language_selector], 
            outputs=[score_interview_button, interview_status, interview_chat_interface.chatbot, interview_chat_interface.chatbot]
        )
        
        score_interview_button.click(
            fn=lambda chat_data, pos, jd, lang: analyze_interview_feedback(chat_data['value'], pos, jd, lang),
            inputs=[interview_chat_interface.chatbot, target_pos_interview, jd_interview, language_selector],
            outputs=interview_feedback_output
        )

# Final step to launch the demo using the best-practice conditional block
if __name__ == "__main__":
    demo.launch()